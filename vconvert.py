import sys
import pdfkit
import docx

# from docx import Document
# from io import BytesIO
# from reportlab.pdfgen import canvas

# def word_to_pdf(word_file, pdf_file):
#     document = Document(word_file)
#     stream = BytesIO()
#     canvas_obj = canvas.Canvas(stream)
#     for paragraph in document.paragraphs:
#         canvas_obj.drawString(100, 750, paragraph.text)
#     canvas_obj.save()
#     stream.seek(0)
#     with open(pdf_file, 'wb') as f:
#         f.write(stream.read())

# if __name__ == '__main__':
#     word_file = sys.argv[1]
#     print(word_file)
#     pdf_file = sys.argv[2]
#     print(pdf_file)
#     word_to_pdf(word_file, pdf_file)

def docx_to_pdf(docx_filename, pdf_filename):
    doc = docx.Document(docx_filename)
    text = ""
    images = []
    for para in doc.paragraphs:
        text += para.text
    for i, pic in enumerate(doc.inline_shapes):
        try:
            pic.save("image" + str(i) + ".png")
            images.append("image" + str(i) + ".png")
        except Exception as e:
            pass
    pdfkit.from_string(text, pdf_filename, options={"images": images})


if __name__ == '__main__':
    word_file = sys.argv[1]
    pdf_file = sys.argv[2]
    docx_to_pdf(word_file, pdf_file)